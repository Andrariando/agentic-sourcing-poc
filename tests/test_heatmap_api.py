"""
API smoke tests for heatmap KPI/KLI, intake meta, approve idempotency, metrics dashboard.
Run from repo root: pytest tests/test_heatmap_api.py -q
"""
import json

import pytest
from fastapi.testclient import TestClient

from backend.main import app


@pytest.fixture
def client():
    with TestClient(app) as c:
        yield c


def test_opportunities_enrich_default_has_kli_metrics(client: TestClient):
    r = client.get("/api/heatmap/opportunities")
    assert r.status_code == 200
    opps = r.json()["opportunities"]
    assert opps
    o0 = opps[0]
    assert "kli_metrics" in o0
    assert "data_quality_warnings" in o0
    assert "feedback_rows" in o0["kli_metrics"]


def test_opportunities_enrich_false_strips_enrichment(client: TestClient):
    r = client.get("/api/heatmap/opportunities?enrich=false")
    assert r.status_code == 200
    o0 = r.json()["opportunities"][0]
    assert "kli_metrics" not in o0
    assert "data_quality_warnings" not in o0


def test_intake_categories_has_card_fingerprint(client: TestClient):
    r = client.get("/api/heatmap/intake/categories")
    assert r.status_code == 200
    data = r.json()
    assert "categories" in data
    assert "category_cards_meta" in data
    assert data["category_cards_meta"].get("sha256")


def test_intake_categories_lists_real_categories_not_file_docs(client: TestClient):
    r = client.get("/api/heatmap/intake/categories")
    assert r.status_code == 200
    cats = r.json()["categories"]
    assert "_documentation" not in cats
    assert all(not str(c).startswith("_") for c in cats)
    assert "IT Infrastructure" in cats


def test_category_scoring_mix_renormalizes_renewal_block():
    from backend.heatmap.category_scoring_mix import apply_category_scoring_overlay
    from backend.heatmap.services.learned_weights import DEFAULT_WEIGHTS_FLAT, normalize_full

    base = normalize_full(dict(DEFAULT_WEIGHTS_FLAT))
    card = {
        "scoring_mix": {
            "existing_contract_renewal": {
                "supplier_risk_RSS": 0.5,
                "expiry_urgency_EUS": 0.2,
                "financial_impact_FIS": 0.1,
                "spend_concentration_SCS": 0.1,
                "strategic_alignment_SAS": 0.1,
            }
        }
    }
    out = apply_category_scoring_overlay(base, card)
    renew_sum = sum(
        out[k] for k in ("w_eus", "w_fis", "w_rss", "w_scs", "w_sas_contract")
    )
    assert renew_sum == pytest.approx(1.0, abs=0.02)
    assert out["w_rss"] > DEFAULT_WEIGHTS_FLAT["w_rss"]


def test_validate_scoring_mix_rejects_unknown_weight_key():
    from backend.heatmap.category_scoring_mix import validate_scoring_mix_for_patch

    with pytest.raises(ValueError, match="Unknown weight key"):
        validate_scoring_mix_for_patch(
            {"new_sourcing_request": {"implementation_typo_IUS": 1.0}}
        )


def test_metrics_dashboard_shape(client: TestClient):
    r = client.get("/api/heatmap/metrics/dashboard")
    assert r.status_code == 200
    d = r.json()
    assert "feedback_rows_total" in d
    assert "opportunities_total" in d
    assert "tier_counts" in d
    assert isinstance(d["feedback_rows_total"], int)


def test_run_status_includes_last_duration_field(client: TestClient):
    r = client.get("/api/heatmap/run/status")
    assert r.status_code == 200
    data = r.json()
    assert "last_duration_sec" in data


def test_approve_idempotent_flags(client: TestClient):
    r = client.get("/api/heatmap/opportunities?enrich=false")
    assert r.status_code == 200
    opps = [o for o in r.json()["opportunities"] if o.get("status") == "Pending"]
    if not opps:
        pytest.skip("no pending opportunities to approve")
    oid = opps[0]["id"]
    a1 = client.post(
        "/api/heatmap/approve",
        json={"opportunity_ids": [oid], "approver_id": "pytest"},
    )
    assert a1.status_code == 200
    j1 = a1.json()
    assert "already_linked" in j1
    assert j1["already_linked"].get(str(oid)) is False
    a2 = client.post(
        "/api/heatmap/approve",
        json={"opportunity_ids": [oid], "approver_id": "pytest"},
    )
    assert a2.status_code == 200
    j2 = a2.json()
    assert j2["already_linked"].get(str(oid)) is True


def test_data_quality_warnings_helper():
    from backend.heatmap.services.data_quality import warnings_for_opportunity

    w = warnings_for_opportunity(
        {"category": "UnknownCat", "source": "intake", "supplier_name": None},
        category_cards_keys=["IT Infrastructure"],
    )
    assert any("category" in x.lower() for x in w)


def test_opportunity_not_pursue_and_filter_exclusion(client: TestClient):
    r = client.get("/api/heatmap/opportunities?enrich=false")
    assert r.status_code == 200
    opps = r.json()["opportunities"]
    assert opps
    oid = opps[0]["id"]

    u = client.post(
        "/api/heatmap/opportunities/disposition",
        json={
            "opportunity_id": oid,
            "disposition": "not_pursuing",
            "not_pursue_reason_code": "strategy_change",
            "comment_text": "No longer in scope",
            "updated_by": "pytest",
        },
    )
    assert u.status_code == 200
    updated = u.json()["opportunity"]
    assert updated["disposition"] == "not_pursuing"
    assert updated["not_pursue_reason_code"] == "strategy_change"

    f = client.get("/api/heatmap/opportunities?enrich=false&include_not_pursuing=false")
    assert f.status_code == 200
    ids = {o["id"] for o in f.json()["opportunities"]}
    assert oid not in ids


def test_case_cancel_with_reason(client: TestClient):
    c = client.post(
        "/api/cases",
        json={"category_id": "IT Infrastructure", "trigger_source": "pytest"},
    )
    assert c.status_code == 200
    case_id = c.json()["case_id"]

    k = client.post(
        f"/api/cases/{case_id}/cancel",
        json={
            "reason_code": "budget_cut",
            "reason_text": "Stopped after budget re-baseline",
            "cancelled_by": "pytest-user",
        },
    )
    assert k.status_code == 200
    assert k.json()["status"] == "Cancelled"

    g = client.get(f"/api/cases/{case_id}")
    assert g.status_code == 200
    assert g.json()["status"] == "Cancelled"


def test_category_cards_extract_from_unstructured_text(client: TestClient):
    raw_text = """
    Default supplier status: allowed
    Category Strategy SAS = 8.5
    Azure: preferred
    LegacyVendor - non preferred
    """
    r = client.post(
        "/api/heatmap/category-cards/extract",
        json={"category": "IT Infrastructure", "raw_text": raw_text},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["used_llm"] is False
    assert data["category"] == "IT Infrastructure"
    patch = data["proposed_patch"]
    assert patch["default_preferred_status"] == "allowed"
    assert patch["category_strategy_sas"] == pytest.approx(8.5)
    assert patch["supplier_preferred_status"]["Azure"] == "preferred"
    assert patch["supplier_preferred_status"]["LegacyVendor"] == "nonpreferred"


def test_category_cards_extract_requires_meaningful_text(client: TestClient):
    r = client.post(
        "/api/heatmap/category-cards/extract",
        json={"category": "IT Infrastructure", "raw_text": "too short"},
    )
    assert r.status_code == 422


def test_category_cards_extract_upload_multipart(client: TestClient):
    body = b"""Default supplier status: allowed
Category Strategy SAS = 8.0
WidgetCo: preferred
"""
    r = client.post(
        "/api/heatmap/category-cards/extract-upload",
        data={"category": "Software"},
        files={"file": ("policy.txt", body, "text/plain")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data.get("filename") == "policy.txt"
    assert data.get("source_format") == ".txt"
    assert data["proposed_patch"]["default_preferred_status"] == "allowed"
    assert data["proposed_patch"]["supplier_preferred_status"]["WidgetCo"] == "preferred"


def test_upload_plain_text_extracts_docx():
    pytest.importorskip("docx", reason="python-docx not installed")
    import io

    from docx import Document

    from backend.heatmap.services.upload_plain_text import upload_bytes_to_plain_text

    buf = io.BytesIO()
    d = Document()
    d.add_paragraph("Default supplier status: allowed")
    d.add_paragraph("Line for extract")
    d.save(buf)
    text = upload_bytes_to_plain_text(buf.getvalue(), "guidance.docx")
    assert "allowed" in text
    assert "extract" in text


def test_category_cards_extract_upload_docx(client: TestClient):
    pytest.importorskip("docx", reason="python-docx not installed")
    import io

    from docx import Document

    buf = io.BytesIO()
    d = Document()
    d.add_paragraph("Default supplier status: allowed")
    d.add_paragraph("Category Strategy SAS = 8.0")
    d.add_paragraph("WidgetCo: preferred")
    d.save(buf)
    r = client.post(
        "/api/heatmap/category-cards/extract-upload",
        data={"category": "Software"},
        files={
            "file": (
                "policy.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200
    data = r.json()
    assert data.get("filename") == "policy.docx"
    assert data.get("source_format") == ".docx"
    assert data["proposed_patch"]["default_preferred_status"] == "allowed"
    assert data["proposed_patch"]["supplier_preferred_status"]["WidgetCo"] == "preferred"


def test_apply_category_cards_patch_merge_tmp_path(tmp_path):
    from backend.heatmap.services.category_cards_store import apply_category_cards_patch

    p = tmp_path / "category_cards.json"
    p.write_text(
        json.dumps(
            {
                "Software": {
                    "default_preferred_status": "allowed",
                    "category_strategy_sas": 7.5,
                    "supplier_preferred_status": {"NexusSoft LLC": "preferred"},
                }
            }
        ),
        encoding="utf-8",
    )
    out = apply_category_cards_patch(
        "Software",
        {"category_strategy_sas": 8.0, "supplier_preferred_status": {"NewCo": "allowed"}},
        cards_path=p,
    )
    assert out["merged_card"]["category_strategy_sas"] == 8.0
    assert out["merged_card"]["supplier_preferred_status"]["NexusSoft LLC"] == "preferred"
    assert out["merged_card"]["supplier_preferred_status"]["NewCo"] == "allowed"
    reloaded = json.loads(p.read_text(encoding="utf-8"))
    assert reloaded["Software"]["category_strategy_sas"] == 8.0


def test_system1_preview_includes_scoring_readiness(client: TestClient):
    csv_bytes = (
        b"row_type,category,supplier_name,estimated_spend_usd,contract_id,months_to_expiry\n"
        b"renewal,IT Infrastructure,VendorA,1200000,CNT-1,6\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("rows.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["candidates"]
    row = data["candidates"][0]
    assert "score_components" in row
    assert "computed_total_score" in row
    assert "readiness_status" in row


def test_system1_approve_requires_warning_ack(client: TestClient):
    # Missing months_to_expiry makes at least one component defaulted => warning row.
    csv_bytes = (
        b"row_type,category,supplier_name,estimated_spend_usd,contract_id\n"
        b"renewal,IT Infrastructure,VendorB,900000,CNT-2\n"
    )
    p = client.post(
        "/api/system1/upload/preview",
        files={"files": ("rows.csv", csv_bytes, "text/csv")},
    )
    assert p.status_code == 200
    payload = p.json()
    rid = payload["candidates"][0]["row_id"]
    job_id = payload["job_id"]

    fail = client.post(
        "/api/system1/upload/approve",
        json={"job_id": job_id, "approved_row_ids": [rid], "approver_id": "pytest"},
    )
    assert fail.status_code == 400

    ok = client.post(
        "/api/system1/upload/approve",
        json={
            "job_id": job_id,
            "approved_row_ids": [rid],
            "approver_id": "pytest",
            "acknowledge_warning_row_ids": [rid],
        },
    )
    assert ok.status_code == 200
    assert ok.json()["success"] is True


def test_system1_templates_endpoint(client: TestClient):
    r = client.get("/api/system1/upload/templates")
    assert r.status_code == 200
    templates = r.json().get("templates", [])
    assert templates
    names = {t["name"] for t in templates}
    assert "renewals_template.csv" in names
    assert "new_business_template.csv" in names


def test_system1_preview_sponsor_style_headers(client: TestClient):
    """Excel/CSV exports often use vendor / commodity / TCV style columns."""
    csv_bytes = (
        b"Opportunity Type,Commodity,Vendor Name,Agreement #,TCV,Expiration Date\n"
        b"renewal,Cloud Services,Acme Corp,AGR-99,\"$1,200,000\",2026-11-30\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("sponsor_export.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["candidates"]
    row = data["candidates"][0]
    assert row["supplier_name"] == "Acme Corp"
    assert row["category"] == "Cloud Services"
    assert row["contract_id"] == "AGR-99"
    assert float(row["estimated_spend_usd"]) == 1_200_000.0
    assert row["row_type"] == "renewal"


def test_system1_preview_capstone_contract_export_headers(client: TestClient):
    """J&J-style contract workbook columns (Supplier, Value in USD, Master Agreement, Expiry)."""
    csv_bytes = (
        b"Supplier,Category,Sub Category,Expiry Date,Value in USD,Master Agreement Reference Number,Contract Name\n"
        b"RENTACOM INC,IT Infrastructure,zIT - Other - DO NOT USE,2016-06-30,2800,C2015023508,Dermatology Work Order\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("contract_export.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    row = r.json()["candidates"][0]
    assert row["supplier_name"] == "RENTACOM INC"
    assert row["category"] == "IT Infrastructure"
    assert row["subcategory"] == "zIT - Other - DO NOT USE"
    assert row["contract_id"] == "C2015023508"
    assert float(row["estimated_spend_usd"]) == 2800.0
    assert "Dermatology" in (row.get("request_title") or "")


def test_system1_preview_capstone_spend_export_headers(client: TestClient):
    """IT Infra Spend grain: booked hierarchy + SMD supplier + invoice amount."""
    csv_bytes = (
        b"key_spend_id,Booked category,Booked subcategory,Commodity,SMD cleansed name,Invoice amount (USD)\n"
        b"k1,IT Infrastructure,IT Infrastructure - Mobile,Mobile Device Management,DISCORP NV,15.76\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("spend_lines.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    row = r.json()["candidates"][0]
    assert row["supplier_name"] == "DISCORP NV"
    assert row["category"] == "IT Infrastructure"
    assert "Mobile" in (row.get("subcategory") or "")
    assert float(row["estimated_spend_usd"]) == 15.76
    assert row.get("request_title")


def test_system1_scan_bundle_fuses_contract_spend_metrics(client: TestClient):
    contract_csv = (
        b"Supplier,Category,Sub Category,Expiry Date,Master Agreement Reference Number,Contract Name\n"
        b"ACME CORP,IT Infrastructure,Cloud Hosting,2027-06-30,CNT-ACME-1,Cloud Renewal\n"
    )
    spend_csv = (
        b"SMD cleansed name,Booked category,Booked subcategory,Invoice amount (USD),key_spend_id\n"
        b"ACME CORP,IT Infrastructure,Cloud Hosting,1000000,sp-1\n"
    )
    metrics_csv = (
        b"Supplier,Category,Subcategory,Supplier Risk Score,BPRA Vendor Status\n"
        b"ACME CORP,IT Infrastructure,Cloud Hosting,7.2,preferred\n"
    )
    r = client.post(
        "/api/system1/upload/scan-bundle",
        files=[
            ("files", ("contracts.csv", contract_csv, "text/csv")),
            ("files", ("spend.csv", spend_csv, "text/csv")),
            ("files", ("metrics.csv", metrics_csv, "text/csv")),
        ],
    )
    assert r.status_code == 200
    data = r.json()
    assert data["total_candidates"] >= 1
    rows = data["candidates"]
    # Contract-led fusion should produce a renewal row for ACME.
    acme = [x for x in rows if (x.get("supplier_name") or "").upper() == "ACME CORP"]
    assert acme
    row = acme[0]
    assert row["row_type"] == "renewal"
    assert row["contract_id"] == "CNT-ACME-1"
    assert float(row["estimated_spend_usd"]) > 0
    assert row.get("computed_total_score") is not None


def test_system1_scan_bundle_top_n_and_completeness_analysis(client: TestClient):
    contract_csv = (
        b"Supplier,Category,Sub Category,Expiry Date,Master Agreement Reference Number,Contract Name,Value in USD\n"
        b"ACME CORP,IT Infrastructure,Cloud Hosting,2027-06-30,CNT-ACME-1,Cloud Renewal,1000000\n"
        b"BETA LLC,IT Infrastructure,Service Desk,2028-01-31,CNT-BETA-1,Service Renewal,500000\n"
    )
    spend_csv = (
        b"SMD cleansed name,Booked category,Booked subcategory,(P) Payment amount (USD),key_spend_id\n"
        b"ACME CORP,IT Infrastructure,Cloud Hosting,1000000,sp-1\n"
        b"BETA LLC,IT Infrastructure,Service Desk,100,sp-2\n"
    )
    metrics_csv = (
        b"Supplier,Category,Subcategory,Supplier Risk Score,BPRA Vendor Status\n"
        b"ACME CORP,IT Infrastructure,Cloud Hosting,7.2,preferred\n"
        b"BETA LLC,IT Infrastructure,Service Desk,4.9,allowed\n"
    )
    r = client.post(
        "/api/system1/upload/scan-bundle",
        data={"top_n": "1", "rank_by": "score"},
        files=[
            ("files", ("contracts.csv", contract_csv, "text/csv")),
            ("files", ("spend.csv", spend_csv, "text/csv")),
            ("files", ("metrics.csv", metrics_csv, "text/csv")),
        ],
    )
    assert r.status_code == 200
    data = r.json()
    assert data["total_candidates"] >= 2
    assert len(data["candidates"]) == 1
    assert "analysis" in data
    assert data["analysis"]["top_n_applied"] == 1
    assert data["analysis"]["rank_by_applied"] == "score"
    assert "imputation_action_candidates" in data["analysis"]
    assert data["analysis"].get("execution_trace") == [
        "fuse_bundle_rows",
        "score_and_enrich_candidates",
        "summarize_completeness",
        "prioritize_top_n",
    ]
    c0 = data["candidates"][0]
    assert "completeness_score" in c0
    assert "suggested_actions" in c0


def test_system1_scan_bundle_preserves_preferred_status_from_contract_rows(client: TestClient):
    contract_csv = (
        b"row_type,category,subcategory,supplier_name,contract_id,contract_end_date,months_to_expiry,estimated_spend_usd,preferred_supplier_status\n"
        b"renewal,IT Infrastructure,Cloud Hosting,ACME CORP,CNT-ACME-2,2027-06-30,12,1000000,preferred\n"
    )
    r = client.post(
        "/api/system1/upload/scan-bundle",
        files=[("files", ("complete_dataset_template_1.csv", contract_csv, "text/csv"))],
    )
    assert r.status_code == 200
    data = r.json()
    assert data["candidates"]
    row = data["candidates"][0]
    assert row["preferred_supplier_status"] == "preferred"
    assert row["row_type"] == "renewal"
    # With preferred status present, RSS should be derived (not defaulted) and warning should not mention fallback.
    warns = " ".join(row.get("warnings") or []).lower()
    assert "rss_score used fallback default" not in warns


def test_system1_preview_rank_by_score_with_top_n(client: TestClient):
    csv_bytes = (
        b"row_type,category,supplier_name,estimated_spend_usd,contract_id,months_to_expiry\n"
        b"renewal,IT Infrastructure,VendorA,1200000,CNT-1,6\n"
        b"renewal,IT Infrastructure,VendorB,1000,CNT-2,2\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        data={"top_n": "1", "rank_by": "score"},
        files={"files": ("rows.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    data = r.json()
    assert len(data["candidates"]) == 1
    assert data["analysis"]["top_n_applied"] == 1
    assert data["analysis"]["rank_by_applied"] == "score"


def test_system1_preview_uses_learned_weight_mix(client: TestClient):
    wr = client.put(
        "/api/heatmap/scoring-weights",
        json={
            "weights": {
                "w_eus": 0.70,
                "w_fis": 0.10,
                "w_rss": 0.10,
                "w_scs": 0.05,
                "w_sas_contract": 0.05,
            }
        },
    )
    assert wr.status_code == 200
    csv_bytes = (
        b"row_type,category,supplier_name,estimated_spend_usd,contract_id,months_to_expiry\n"
        b"renewal,ZZZ Unmapped Category,VendorA,1200000,CNT-1,2\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("rows.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    row = r.json()["candidates"][0]
    wu = row.get("weights_used") or {}
    assert abs(float(wu.get("w_eus", 0.0)) - 0.70) < 0.02
    assert abs(float(wu.get("w_fis", 0.0)) - 0.10) < 0.02


def test_system1_preview_preserves_explicit_new_business_even_with_contract_id(client: TestClient):
    csv_bytes = (
        b"row_type,category,supplier_name,estimated_spend_usd,contract_id,request_title,implementation_timeline_months\n"
        b"new_business,IT Infrastructure,,1200000,CNT-NEW-1,ERP Request,6\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("rows.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    row = r.json()["candidates"][0]
    assert row["row_type"] == "new_business"


def test_system1_preview_new_business_without_supplier_can_still_be_complete(client: TestClient):
    csv_bytes = (
        b"row_type,category,supplier_name,estimated_spend_usd,request_title,implementation_timeline_months\n"
        b"new_business,IT Infrastructure,,1200000,ERP Request,6\n"
    )
    r = client.post(
        "/api/system1/upload/preview",
        files={"files": ("rows.csv", csv_bytes, "text/csv")},
    )
    assert r.status_code == 200
    row = r.json()["candidates"][0]
    assert row["row_type"] == "new_business"
    assert row.get("completeness_score") == 100.0
    assert "supplier_name" not in (row.get("missing_critical_fields") or [])


def test_system1_preview_erp_adapter_normalizes_rows(client: TestClient):
    payload = {
        "source_system": "sap_s4",
        "mapping_profile": "erp_generic",
        "rows": [
            {
                "opportunity_type": "new_business",
                "vendor_name": "Acme Corp",
                "request_name": "Network Segmentation Rollout",
                "amount_usd": "450000",
                "implementation_months": 5,
                "category": "IT Infrastructure",
            },
            {
                "opportunity_type": "renewal",
                "vendor_name": "Blue Systems",
                "agreement_number": "AGR-1001",
                "amount_usd": "900000",
                "months_remaining": 4,
                "category": "IT Infrastructure",
            },
        ],
    }
    r = client.post("/api/system1/upload/preview-erp", json=payload)
    assert r.status_code == 200
    data = r.json()
    assert len(data["candidates"]) == 2
    types = sorted([c["row_type"] for c in data["candidates"]])
    assert types == ["new_business", "renewal"]
    assert "analysis" in data
    assert "ingestion_diagnostics" in data["analysis"]


def test_system1_scan_bundle_rejects_non_structured_only(client: TestClient):
    txt_bytes = b"This is guidance text only."
    r = client.post(
        "/api/system1/upload/scan-bundle",
        files=[("files", ("guidance.txt", txt_bytes, "text/plain"))],
    )
    assert r.status_code == 400


def test_scoring_config_endpoints_crud_validate_publish(client: TestClient):
    active = client.get("/api/heatmap/scoring-config/active")
    assert active.status_code == 200
    active_json = active.json()
    assert active_json["status"] == "active"
    assert isinstance(active_json.get("config"), dict)

    draft_payload = {
        "title": "Pytest scoring config draft",
        "created_by": "pytest",
        "config": active_json["config"],
    }
    created = client.post("/api/heatmap/scoring-config/draft", json=draft_payload)
    assert created.status_code == 200
    draft = created.json()
    assert draft["status"] == "draft"
    draft_id = int(draft["id"])

    validated = client.post(f"/api/heatmap/scoring-config/draft/{draft_id}/validate")
    assert validated.status_code == 200
    assert validated.json().get("valid") is True

    published = client.post(f"/api/heatmap/scoring-config/draft/{draft_id}/publish")
    assert published.status_code == 200
    assert published.json()["status"] == "active"

    versions = client.get("/api/heatmap/scoring-config/versions")
    assert versions.status_code == 200
    rows = versions.json()
    assert any(v["status"] == "active" and int(v["id"]) == draft_id for v in rows)


def test_scoring_config_create_draft_rejects_invalid_config(client: TestClient):
    bad = client.post(
        "/api/heatmap/scoring-config/draft",
        json={
            "title": "Bad config",
            "created_by": "pytest",
            "config": {
                "name": "bad",
                "parameters": [
                    {
                        "key": "x",
                        "label": "X",
                        "applies_to": ["renewal"],
                        "input_fields": [],
                        "rule_type": "direct_numeric",
                        "rule_config": {},
                        "default_policy": {"strategy": "constant", "value": 1},
                        "weight_key": "w_not_real",
                    }
                ],
                "formulas": {
                    "renewal": {"weight_keys": ["w_eus"]},
                    "new_business": {"weight_keys": ["w_ius"]},
                },
            },
        },
    )
    assert bad.status_code == 400


def test_scoring_config_publish_syncs_learned_weights(client: TestClient):
    base_cfg = client.get("/api/heatmap/scoring-config/active").json()["config"]
    cfg = json.loads(json.dumps(base_cfg))
    cfg["formulas"]["renewal"]["weight_values"] = {
        "w_eus": 0.55,
        "w_fis": 0.20,
        "w_rss": 0.10,
        "w_scs": 0.10,
        "w_sas_contract": 0.05,
    }
    cfg["formulas"]["new_business"]["weight_values"] = {
        "w_ius": 0.40,
        "w_es": 0.25,
        "w_csis": 0.20,
        "w_sas_new": 0.15,
    }
    created = client.post(
        "/api/heatmap/scoring-config/draft",
        json={"title": "weights sync draft", "created_by": "pytest", "config": cfg},
    )
    assert created.status_code == 200
    draft_id = int(created.json()["id"])
    pub = client.post(f"/api/heatmap/scoring-config/draft/{draft_id}/publish")
    assert pub.status_code == 200

    wr = client.get("/api/heatmap/scoring-weights")
    assert wr.status_code == 200
    w = wr.json()["weights"]
    assert abs(float(w["w_eus"]) - 0.55) < 0.02
    assert abs(float(w["w_ius"]) - 0.40) < 0.02


def test_stage_intake_roundtrip_and_extract(client: TestClient):
    create = client.post(
        "/api/cases",
        json={"category_id": "IT Infrastructure", "trigger_source": "User", "name": "Stage Intake Test"},
    )
    assert create.status_code == 200
    case_id = create.json()["case_id"]

    put = client.put(
        f"/api/cases/{case_id}/stage-intake",
        json={
            "stage": "DTP-01",
            "values": {
                "request_title": "IT Service Desk Renewal",
                "business_unit": "Global Ops",
            },
            "source": "human_form",
            "updated_by": "tester",
        },
    )
    assert put.status_code == 200
    assert put.json()["success"] is True

    get = client.get(f"/api/cases/{case_id}/stage-intake?stage=DTP-01")
    assert get.status_code == 200
    vals = get.json()["values"]
    assert vals.get("request_title") == "IT Service Desk Renewal"
    assert vals.get("business_unit") == "Global Ops"

    extract = client.post(
        f"/api/cases/{case_id}/stage-intake/extract",
        json={
            "stage": "DTP-01",
            "free_text": "request title: APAC Helpdesk refresh; business unit: APAC Operations; estimated value: $1250000",
            "existing_values": vals,
        },
    )
    assert extract.status_code == 200
    proposed = extract.json()["proposed_values"]
    assert proposed.get("business_unit")
    assert proposed.get("estimated_annual_value_usd")


def test_stage_generation_check_preconditions(client: TestClient):
    create = client.post(
        "/api/cases",
        json={"category_id": "IT Infrastructure", "trigger_source": "User", "name": "Generation Check Test"},
    )
    assert create.status_code == 200
    case_id = create.json()["case_id"]

    missing = client.post(
        f"/api/cases/{case_id}/stage-intake/generation-check",
        json={"stage": "DTP-01", "values": {"request_title": "Only title"}},
    )
    assert missing.status_code == 200
    j = missing.json()
    assert j["can_generate"] is False
    assert "business_unit" in j["missing_fields"]

    ok = client.post(
        f"/api/cases/{case_id}/stage-intake/generation-check",
        json={
            "stage": "DTP-01",
            "values": {
                "request_title": "IT Service Desk",
                "business_unit": "Global Ops",
                "scope_summary": "Full desk support",
                "estimated_annual_value_usd": "1000000",
            },
        },
    )
    assert ok.status_code == 200
    assert ok.json()["can_generate"] is True

