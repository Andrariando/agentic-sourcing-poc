"""Unit tests for Word round-trip helpers (no live OpenAI)."""
from io import BytesIO

from docx import Document


def test_extract_text_from_docx_bytes_roundtrip():
    from backend.services.docx_text import extract_text_from_docx_bytes

    buf = BytesIO()
    d = Document()
    d.add_paragraph("Line A")
    d.add_paragraph("Line B")
    d.save(buf)
    text = extract_text_from_docx_bytes(buf.getvalue())
    assert "Line A" in text and "Line B" in text


def test_build_plain_text_docx_bytes_is_docx_zip():
    from backend.services.artifact_document_export import build_plain_text_docx_bytes

    raw = build_plain_text_docx_bytes(
        "RFx draft (working copy)",
        "Hello\n\nSecond paragraph.",
        "CASE-TEST",
        "Test Case",
    )
    assert raw[:2] == b"PK"


def test_format_working_documents_prompt():
    from shared.working_documents_prompt import format_working_documents_for_prompt

    block = format_working_documents_for_prompt(
        {
            "rfx": {
                "plain_text": "Scope: widgets",
                "source_filename": "a.docx",
                "updated_at": "2026-01-01",
                "updated_by": "user_upload",
            }
        }
    )
    assert "RFx" in block or "widgets" in block
