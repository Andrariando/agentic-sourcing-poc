"""
Plain-text extraction for heatmap uploads (category cards, policy extracts).

Supports UTF-8 text and Word .docx (paragraphs + table cells) via python-docx.
"""
from __future__ import annotations

import io
from pathlib import Path


def upload_bytes_to_plain_text(raw: bytes, filename: str) -> str:
    """
    Decode upload bytes to a single string for downstream NLP / deterministic extract.

    Raises:
        ValueError: .docx present but python-docx missing, or no extractable text.
    """
    if not raw:
        return ""

    suffix = Path(filename or "").suffix.lower()
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise ValueError(
                "python-docx is required to read .docx files. Install with: pip install python-docx"
            ) from e

        doc = DocxDocument(io.BytesIO(raw))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)

        text = "\n\n".join(parts).strip()
        if not text:
            raise ValueError("No text could be extracted from the Word document.")
        return text

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="replace")
