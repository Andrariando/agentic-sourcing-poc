"""
Export legacy DTP artifacts to Word (.docx) or PDF for download.

RFx drafts (RFX_DRAFT_PACK) map section structure to headings; other types fall back
to summary text plus JSON content.
"""
from __future__ import annotations

import json
import re
from io import BytesIO
from typing import Any, List, Tuple

from docx import Document
from docx.shared import Pt

from shared.constants import ArtifactType
from shared.schemas import Artifact, ArtifactPack


def _safe_json(obj: Any) -> str:
    try:
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except (TypeError, ValueError):
        return str(obj)


def _ascii_for_pdf(text: str, max_len: int = 120_000) -> str:
    """FPDF core fonts are Latin-1; replace unsupported chars."""
    t = (text or "")[:max_len]
    return t.encode("latin-1", errors="replace").decode("latin-1")


def iter_export_sections(artifact: Artifact) -> List[Tuple[str, str]]:
    """
    Build (heading, body) pairs for document export.
    """
    t = artifact.type or ""
    c = artifact.content or {}

    blocks: List[Tuple[str, str]] = []

    if t == ArtifactType.RFX_DRAFT_PACK.value:
        rfx_type = c.get("rfx_type", "RFx")
        lines = [
            f"Type: {rfx_type}",
            f"Completeness score: {c.get('completeness_score', '—')}%",
            f"Complete: {c.get('is_complete', False)}",
        ]
        if c.get("missing_sections"):
            lines.append("Missing sections: " + ", ".join(str(x) for x in c["missing_sections"]))
        if c.get("incomplete_sections"):
            lines.append("Incomplete sections: " + ", ".join(str(x) for x in c["incomplete_sections"]))
        blocks.append(("Overview", "\n".join(lines)))
        for sec in c.get("sections") or []:
            name = sec.get("section") or "Section"
            body = sec.get("content") or ""
            st = sec.get("status") or ""
            if st:
                body = f"{body}\n\n(Status: {st})"
            blocks.append((name, str(body)))
        return blocks

    if t == ArtifactType.RFX_PATH.value:
        rat = c.get("rationale") or []
        rtxt = "\n".join(f"• {x}" for x in rat) if rat else "(none)"
        blocks.append(("Rationale", rtxt))
        blocks.append(("Missing information", _safe_json(c.get("missing_info"))))
        return blocks

    if t == ArtifactType.RFX_QA_TRACKER.value:
        blocks.append(
            ("Q&A tracker", _safe_json(c.get("tracker") or c.get("questions") or c))
        )
        return blocks

    # Generic: strategy, negotiation, signals, etc.
    if artifact.content_text:
        blocks.append(("Summary", artifact.content_text))
    if c:
        blocks.append(("Structured content", _safe_json(c)))
    elif not artifact.content_text:
        blocks.append(("Content", "(empty)"))
    return blocks


def _sanitize_filename_stem(title: str, fallback: str = "artifact") -> str:
    base = (title or fallback).strip() or fallback
    base = re.sub(r'[<>:"/\\|?*\n\r\t]', "_", base)
    base = re.sub(r"\s+", "_", base).strip("._")
    return base[:100] or fallback


def build_artifact_docx_bytes(artifact: Artifact, case_id: str, case_name: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(case_name or case_id, level=0)
    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_paragraph(
        f"Artifact: {artifact.title} | Type: {artifact.type} | Agent: {artifact.created_by_agent}"
    )
    doc.add_paragraph("")

    for heading, body in iter_export_sections(artifact):
        doc.add_heading(heading, level=1)
        for line in (body or "").split("\n"):
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_artifact_pdf_bytes(artifact: Artifact, case_id: str, case_name: str) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError as e:
        raise RuntimeError("fpdf2 is required for PDF export. Install: pip install fpdf2") from e

    class DocPDF(FPDF):
        def __init__(self) -> None:
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

    pdf = DocPDF()
    pdf.set_margins(12, 12, 12)
    pdf.add_page()
    col_w = getattr(pdf, "epw", pdf.w - pdf.l_margin - pdf.r_margin)

    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(col_w, 8, _ascii_for_pdf(case_name or case_id))
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(col_w, 5, _ascii_for_pdf(f"Case ID: {case_id}"))
    pdf.multi_cell(
        col_w,
        5,
        _ascii_for_pdf(
            f"Artifact: {artifact.title} | Type: {artifact.type} | Agent: {artifact.created_by_agent}"
        ),
    )
    pdf.ln(4)

    for heading, body in iter_export_sections(artifact):
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(col_w, 6, _ascii_for_pdf(heading))
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(col_w, 5, _ascii_for_pdf(body))
        pdf.ln(3)

    out = pdf.output()
    return bytes(out) if not isinstance(out, bytes) else out


def export_filename(artifact: Artifact, ext: str) -> str:
    stem = _sanitize_filename_stem(artifact.title)
    return f"{stem}.{ext.lstrip('.')}"


def build_plain_text_docx_bytes(
    document_title: str, body: str, case_id: str, case_name: str
) -> bytes:
    """
    Build a simple Word document from arbitrary plain text (paragraph breaks preserved).
    Used for exporting user/copilot working copies (RFX / contract slots).
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_heading(case_name or case_id, level=0)
    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_heading(document_title, level=1)
    doc.add_paragraph("")
    for line in (body or "").split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_working_doc_filename(role: str, ext: str) -> str:
    stem = "RFX_draft" if role == "rfx" else "Contract_draft"
    return f"{stem}.{ext.lstrip('.')}"


def export_pack_filename(pack: ArtifactPack, ext: str) -> str:
    agent = _sanitize_filename_stem(pack.agent_name or "artifact_pack", "artifact_pack")
    short_id = (pack.pack_id or "pack").replace(" ", "")[:12]
    return f"{agent}_{short_id}.{ext.lstrip('.')}"


def _pack_markdown_body(pack: ArtifactPack, case_id: str, case_name: str) -> str:
    lines: List[str] = [
        f"# {case_name or case_id}",
        "",
        f"**Case ID:** `{case_id}`",
        "",
        f"## Artifact pack `{pack.pack_id}`",
        "",
        f"- **Agent:** {pack.agent_name or '—'}",
        f"- **Created:** {pack.created_at or '—'}",
        f"- **Artifacts:** {len(pack.artifacts)}",
        "",
    ]
    if pack.next_actions:
        lines.append("### Suggested next actions")
        for a in pack.next_actions:
            lines.append(f"- **{a.label}** — {a.why or ''}".rstrip())
        lines.append("")
    if pack.risks:
        lines.append("### Risks flagged")
        for r in pack.risks:
            lines.append(f"- [{r.severity}] {r.description}")
        lines.append("")

    for art in pack.artifacts:
        lines.append(f"### {art.title}")
        lines.append("")
        lines.append(f"*Type `{art.type}` · Agent `{art.created_by_agent}`*")
        lines.append("")
        for heading, body in iter_export_sections(art):
            lines.append(f"#### {heading}")
            lines.append("")
            lines.append(body or "")
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_artifact_pack_markdown_bytes(pack: ArtifactPack, case_id: str, case_name: str) -> bytes:
    text = _pack_markdown_body(pack, case_id, case_name)
    return text.encode("utf-8")


def build_artifact_pack_docx_bytes(pack: ArtifactPack, case_id: str, case_name: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(case_name or case_id, level=0)
    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_heading(f"Artifact pack {pack.pack_id}", level=1)
    doc.add_paragraph(f"Agent: {pack.agent_name or '—'}")
    doc.add_paragraph(f"Created: {pack.created_at or '—'}")
    doc.add_paragraph("")

    if pack.next_actions:
        doc.add_heading("Suggested next actions", level=2)
        for a in pack.next_actions:
            doc.add_paragraph(f"{a.label} — {a.why or ''}", style="List Bullet")
    if pack.risks:
        doc.add_heading("Risks flagged", level=2)
        for r in pack.risks:
            doc.add_paragraph(f"[{r.severity}] {r.description}", style="List Bullet")

    for art in pack.artifacts:
        doc.add_heading(art.title, level=1)
        meta = doc.add_paragraph()
        r = meta.add_run(f"Type: {art.type} | Agent: {art.created_by_agent}")
        r.italic = True
        for heading, body in iter_export_sections(art):
            doc.add_heading(heading, level=2)
            for line in (body or "").split("\n"):
                doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_artifact_pack_pdf_bytes(pack: ArtifactPack, case_id: str, case_name: str) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError as e:
        raise RuntimeError("fpdf2 is required for PDF export. Install: pip install fpdf2") from e

    class DocPDF(FPDF):
        def __init__(self) -> None:
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

    pdf = DocPDF()
    pdf.set_margins(12, 12, 12)
    pdf.add_page()
    col_w = getattr(pdf, "epw", pdf.w - pdf.l_margin - pdf.r_margin)

    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(col_w, 8, _ascii_for_pdf(case_name or case_id))
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(col_w, 5, _ascii_for_pdf(f"Case ID: {case_id}"))
    pdf.multi_cell(col_w, 5, _ascii_for_pdf(f"Pack: {pack.pack_id} | Agent: {pack.agent_name or '—'}"))
    pdf.ln(3)

    if pack.next_actions:
        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(col_w, 6, _ascii_for_pdf("Suggested next actions"))
        pdf.set_font("Helvetica", "", 9)
        for a in pack.next_actions:
            pdf.multi_cell(col_w, 5, _ascii_for_pdf(f"- {a.label}: {a.why or ''}"))
        pdf.ln(2)

    for art in pack.artifacts:
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(col_w, 6, _ascii_for_pdf(art.title))
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(
            col_w,
            4,
            _ascii_for_pdf(f"Type: {art.type} | Agent: {art.created_by_agent}"),
        )
        for heading, body in iter_export_sections(art):
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(col_w, 5, _ascii_for_pdf(heading))
            pdf.set_font("Helvetica", "", 9)
            pdf.multi_cell(col_w, 5, _ascii_for_pdf(body))
            pdf.ln(2)

    out = pdf.output()
    return bytes(out) if not isinstance(out, bytes) else out
